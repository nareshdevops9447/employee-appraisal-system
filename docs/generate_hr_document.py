"""
Generate EAS Business Logic document for HR team in Word (.docx) and PDF formats.
Run: python docs/generate_hr_document.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour palette ──────────────────────────────────────────────────
BRAND_DARK   = RGBColor(0x1B, 0x26, 0x3B)   # Dark navy
BRAND_PRIMARY = RGBColor(0x2B, 0x57, 0x9A)  # Blue
BRAND_ACCENT  = RGBColor(0x17, 0x84, 0x4E)  # Green
BRAND_LIGHT   = RGBColor(0x5B, 0x5B, 0x5B)  # Grey for body
BRAND_ORANGE  = RGBColor(0xE8, 0x6C, 0x00)  # Highlight
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = "D6E4F0"
LIGHT_GREEN_BG = "D9EAD3"
LIGHT_GREY_BG = "F2F2F2"
LIGHT_ORANGE_BG = "FDE8D0"


def set_cell_shading(cell, color_hex):
    """Set background colour of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows, col_widths=None, header_color=LIGHT_BLUE_BG):
    """Add a nicely formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = BRAND_DARK
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = BRAND_LIGHT
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph('')  # Spacer
    return table


def add_info_box(doc, text, bg_color=LIGHT_GREEN_BG, icon=""):
    """Add a highlighted info box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    p = cell.paragraphs[0]
    run = p.add_run(f"{icon}  {text}" if icon else text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = BRAND_DARK
    run.font.italic = True
    doc.add_paragraph('')


def add_flow_box(doc, title, steps):
    """Add a visual flow with numbered steps."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_PRIMARY

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        # Number circle
        run = p.add_run(f"  {i}  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_ACCENT
        # Step text
        run = p.add_run(step)
        run.font.size = Pt(9.5)
        run.font.color.rgb = BRAND_LIGHT
    doc.add_paragraph('')


def heading(doc, text, level=1):
    """Add styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_DARK if level == 1 else BRAND_PRIMARY
    return h


def body(doc, text, bold_phrases=None):
    """Add body paragraph with optional bold phrases."""
    p = doc.add_paragraph()
    if bold_phrases:
        remaining = text
        for phrase in bold_phrases:
            if phrase in remaining:
                before, _, after = remaining.partition(phrase)
                if before:
                    run = p.add_run(before)
                    run.font.size = Pt(10)
                    run.font.color.rgb = BRAND_LIGHT
                run = p.add_run(phrase)
                run.font.size = Pt(10)
                run.font.color.rgb = BRAND_DARK
                run.bold = True
                remaining = after
        if remaining:
            run = p.add_run(remaining)
            run.font.size = Pt(10)
            run.font.color.rgb = BRAND_LIGHT
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_LIGHT
    return p


def bullet(doc, text, bold_part=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and bold_part in text:
        before, _, after = text.partition(bold_part)
        if before:
            run = p.add_run(before)
            run.font.size = Pt(9.5)
            run.font.color.rgb = BRAND_LIGHT
        run = p.add_run(bold_part)
        run.font.size = Pt(9.5)
        run.font.color.rgb = BRAND_DARK
        run.bold = True
        if after:
            run = p.add_run(after)
            run.font.size = Pt(9.5)
            run.font.color.rgb = BRAND_LIGHT
    else:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = BRAND_LIGHT
    return p


def build_document():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = BRAND_LIGHT

    # ════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Employee Appraisal System')
    run.font.size = Pt(32)
    run.font.color.rgb = BRAND_DARK
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Business Logic & Process Flow')
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_PRIMARY

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('For HR Team Review & Validation')
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_ACCENT

    for _ in range(4):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CONFIDENTIAL')
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_ORANGE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('March 2026  |  Version 1.0')
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_LIGHT

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  PURPOSE
    # ════════════════════════════════════════════════════════════════
    heading(doc, 'Purpose of This Document', level=1)
    body(doc,
         'This document explains how the Employee Appraisal System works in plain, '
         'non-technical language. It is intended for the HR team to review and confirm '
         'whether the logic and process built into the system matches the intended '
         'business requirements.')
    doc.add_paragraph('')
    add_info_box(doc,
                 'Please review each section and confirm whether the logic is correct. '
                 'Use the validation checklist at the end to mark each rule as Correct or Needs Change.',
                 LIGHT_ORANGE_BG, icon="\u26A0")

    # ════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1.  Who Can Do What (Roles)',
        '2.  Appraisal Cycle Setup',
        '3.  Who Is Eligible for Appraisal',
        '4.  The Appraisal Journey (Step by Step)',
        '5.  Goal Setting & Approval',
        '6.  Competency / Attribute Ratings',
        '7.  Peer Feedback',
        '8.  How the Final Score Is Calculated',
        '9.  Calibration (Optional)',
        '10. Employee Acknowledgement & Disputes',
        '11. Appeals (After Completion)',
        '12. Probation Employees (Special Handling)',
        '13. Mid-Year Reviews (Team Transfers)',
        '14. Login & Access',
        '15. Summary of Key Rules & Validation Checklist',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = BRAND_PRIMARY

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 1: ROLES
    # ════════════════════════════════════════════════════════════════
    heading(doc, '1. Who Can Do What (Roles)', level=1)
    body(doc, 'The system has four roles. Each person sees only what is relevant to them.')

    add_styled_table(doc,
        ['Role', 'What They Can Do'],
        [
            ['Employee', 'Set their own goals, do self-assessment, rate themselves on competencies, acknowledge their final appraisal, raise appeals or disputes'],
            ['Manager', 'Approve or reject team goals, review & rate team members, submit manager assessments, request peer feedback for their team'],
            ['HR Admin', 'Create and manage appraisal cycles, set up competency templates, run calibration sessions, handle appeals, view all appraisals across the organisation'],
            ['Super Admin', 'Everything HR Admin can do plus system-level settings and user management'],
        ],
        col_widths=[4, 14]
    )

    heading(doc, 'Key Access Rules', level=2)
    bullet(doc, 'An employee can only see their own appraisal', 'their own appraisal')
    bullet(doc, 'A manager can see appraisals of their direct reports only', 'direct reports only')
    bullet(doc, 'HR can see all appraisals across the entire organisation', 'entire organisation')
    bullet(doc, "A person's role can change automatically \u2014 for example, if Azure AD shows someone now has direct reports, they become a Manager",
           'automatically')
    doc.add_paragraph('')

    # ════════════════════════════════════════════════════════════════
    #  SECTION 2: CYCLE SETUP
    # ════════════════════════════════════════════════════════════════
    heading(doc, '2. Appraisal Cycle Setup', level=1)
    body(doc, 'An appraisal cycle is the time period during which appraisals happen (e.g., "Annual Review 2026"). HR configures the following settings:')

    add_styled_table(doc,
        ['Setting', 'What It Means'],
        [
            ['Cycle Name', 'e.g., "Annual Appraisal 2026"'],
            ['Cycle Type', 'Annual, Probation, or Mid-Year'],
            ['Start & End Dates', 'The period the cycle covers'],
            ['Self-Assessment Deadline', 'Last date for employees to complete self-ratings'],
            ['Manager Review Deadline', 'Last date for managers to submit their reviews'],
            ['Eligibility Cutoff Date', 'Employees who joined AFTER this date are deferred to the probation cycle'],
            ['Minimum Service Months', 'How many months someone must have worked to be eligible'],
            ['Include Probation Employees?', 'Yes/No \u2014 should employees still on probation be included?'],
            ['Allow Prorated Evaluation?', 'Yes/No \u2014 for mid-cycle joiners, should they get a prorated review?'],
            ['Requires Calibration?', 'Yes/No \u2014 should there be an HR calibration step before the employee sees their rating?'],
            ['Score Weights', 'How much each component counts toward the final score (must add up to 100%)'],
        ],
        col_widths=[5.5, 12.5]
    )

    heading(doc, 'Score Weight Example', level=2)
    add_styled_table(doc,
        ['Component', 'Weight', 'Meaning'],
        [
            ['Goals', '70%', 'How well they achieved their objectives'],
            ['Competencies', '30%', 'How they demonstrated expected behaviours'],
            ['Peer Feedback', '0%', 'Not used this cycle'],
            ['TOTAL', '100%', 'Must always add up to 100%'],
        ],
        col_widths=[4.5, 3, 10.5],
        header_color=LIGHT_GREEN_BG
    )

    heading(doc, 'Important Rules', level=2)
    bullet(doc, 'Only ONE cycle of each type can be active at a time (e.g., you cannot have two annual cycles running simultaneously)',
           'ONE cycle of each type')
    bullet(doc, 'When HR activates an annual cycle, the system automatically creates a companion "Probation Goals" cycle for deferred employees',
           'automatically creates')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 3: ELIGIBILITY
    # ════════════════════════════════════════════════════════════════
    heading(doc, '3. Who Is Eligible for Appraisal', level=1)
    body(doc, 'When a cycle is activated, the system checks every active employee and decides if they qualify.')

    heading(doc, 'Annual Cycle \u2014 Eligibility Checklist', level=2)
    body(doc, 'The system goes through each check in order. The first one that fails determines the outcome:')

    add_styled_table(doc,
        ['Check', 'If Fails...', 'Result'],
        [
            ['Have they completed their 3-month probation?', 'Still within probation', 'Deferred to Probation Cycle'],
            ['Does their probation end before the cycle ends?', 'Probation extends beyond cycle', 'Deferred to Probation Cycle'],
            ['Did they join before the eligibility cutoff date?', 'Joined after cutoff', 'Deferred to Probation Cycle'],
            ['Have they worked the minimum required months?', 'Insufficient service', 'Not Eligible'],
            ['Is the cycle set to include probation employees?', 'Probation excluded', 'Not Eligible'],
            ['Is prorated evaluation allowed? (for mid-cycle joiners)', 'Proration disabled', 'Not Eligible'],
            ['Did they join after the cycle started? (with proration on)', '\u2014', 'Eligible, but PRORATED'],
            ['All checks passed', '\u2014', 'Fully Eligible'],
        ],
        col_widths=[6, 5, 5]
    )

    heading(doc, 'What Happens to Deferred Employees', level=2)
    bullet(doc, 'They are NOT ignored \u2014 they are automatically placed into the Probation Cycle instead',
           'NOT ignored')
    bullet(doc, 'The system creates this probation cycle automatically when the annual cycle is activated',
           'automatically')

    heading(doc, 'What "Prorated" Means', level=2)
    bullet(doc, 'The employee joined mid-cycle, so their evaluation period is shorter')
    bullet(doc, 'This flag is visible to HR and the manager for context')
    bullet(doc, 'The scores are still calculated the same way \u2014 it is informational only')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 4: APPRAISAL JOURNEY
    # ════════════════════════════════════════════════════════════════
    heading(doc, '4. The Appraisal Journey (Step by Step)', level=1)
    body(doc, 'Every appraisal moves through these stages in order. It cannot skip steps or go backwards (except for goals being revised).')

    # Visual flow as a table
    flow_table = doc.add_table(rows=1, cols=4)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stages = [
        ('STAGE 1', 'Goal Setting', LIGHT_BLUE_BG),
        ('STAGE 2', 'Self-Assessment', LIGHT_GREEN_BG),
        ('STAGE 3', 'Manager Review', LIGHT_ORANGE_BG),
        ('STAGE 4', 'Sign-Off', LIGHT_BLUE_BG),
    ]
    for i, (num, name, bg) in enumerate(stages):
        cell = flow_table.rows[0].cells[i]
        set_cell_shading(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{num}\n{name}')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_DARK

    # Arrow description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Goal Setting  \u2192  Self-Assessment  \u2192  Manager Review  \u2192  (Calibration)  \u2192  Acknowledgement  \u2192  Completed')
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_PRIMARY
    run.italic = True
    doc.add_paragraph('')

    # Stage 1
    heading(doc, 'Stage 1: Goal Setting', level=2)
    add_styled_table(doc,
        ['What Happens', 'Who Does It'],
        [
            ['Performance goals are created (exactly 5 required)', 'Employee or Manager'],
            ['Each goal gets a weight (all weights must add up to 100%)', 'Employee or Manager'],
            ['Goals are submitted for approval', 'Employee submits \u2192 Manager approves'],
            ['Manager can approve or reject each goal', 'Manager'],
            ['If rejected, employee revises and resubmits', 'Employee'],
            ['All 5 performance goals must be approved to move forward', '\u2014'],
            ['Development goals can also be set (optional, max 3)', 'Employee'],
        ],
        col_widths=[11, 7]
    )
    add_info_box(doc, 'Development goals are for learning/growth only. They do NOT count toward the final score and are not required to move forward.', LIGHT_GREEN_BG)

    # Stage 2
    heading(doc, 'Stage 2: Self-Assessment', level=2)
    add_styled_table(doc,
        ['What Happens', 'Who Does It'],
        [
            ['Employee rates themselves on each goal (1\u20135 scale)', 'Employee'],
            ['Employee rates themselves on each competency (1\u20135 scale)', 'Employee'],
            ['Employee can add comments for each rating', 'Employee'],
            ['Employee submits their self-assessment', 'Employee'],
            ['Deadline is enforced (set by HR in the cycle)', 'System'],
        ],
        col_widths=[11, 7]
    )

    # Stage 3
    heading(doc, 'Stage 3: Manager Review', level=2)
    add_styled_table(doc,
        ['What Happens', 'Who Does It'],
        [
            ["Manager sees employee's self-ratings and comments", 'Manager'],
            ['Manager provides their own rating for each goal (1\u20135)', 'Manager'],
            ['Manager provides their own rating for each competency (1\u20135)', 'Manager'],
            ['Manager writes narrative: strengths, development areas, overall comments', 'Manager'],
            ['Manager submits their assessment', 'Manager'],
            ['System automatically calculates the final score', 'System'],
            ['Deadline is enforced (set by HR in the cycle)', 'System'],
        ],
        col_widths=[11, 7]
    )

    # Stage 3.5
    heading(doc, 'Stage 3.5: Calibration (Only If Turned On)', level=2)
    add_styled_table(doc,
        ['What Happens', 'Who Does It'],
        [
            ['After manager submits, the appraisal goes to HR instead of the employee', 'System'],
            ['HR reviews scores across the organisation for fairness', 'HR'],
            ['HR can adjust the overall rating (1\u20135)', 'HR'],
            ['HR can add calibration notes', 'HR'],
            ['HR approves and pushes to employee for acknowledgement', 'HR'],
        ],
        col_widths=[11, 7]
    )

    # Stage 4
    heading(doc, 'Stage 4: Employee Acknowledgement', level=2)
    add_styled_table(doc,
        ['What Happens', 'Who Does It'],
        [
            ['Employee sees their final rating and all manager feedback', 'Employee'],
            ['Employee acknowledges the appraisal', 'Employee'],
            ['Employee can add comments', 'Employee'],
            ['Employee can flag a dispute (disagree with the rating)', 'Employee'],
            ['If disputed, the manager is notified', 'System'],
            ['The appraisal is marked as Completed', 'System'],
        ],
        col_widths=[11, 7]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 5: GOAL SETTING
    # ════════════════════════════════════════════════════════════════
    heading(doc, '5. Goal Setting & Approval', level=1)

    heading(doc, 'Goal Types', level=2)
    add_styled_table(doc,
        ['Type', 'Count', 'Required?', 'Affects Score?'],
        [
            ['Performance', 'Exactly 5', 'Yes \u2014 cannot proceed without them', 'Yes \u2014 weighted into final score'],
            ['Development', 'Up to 3', 'No \u2014 completely optional', 'No \u2014 for personal growth only'],
        ],
        col_widths=[4, 3, 5.5, 5.5]
    )

    heading(doc, 'Performance Goal Requirements', level=2)
    bullet(doc, 'Each goal has a weight (a percentage showing its importance)')
    bullet(doc, 'All 5 weights must add up to exactly 100%', '100%')
    bullet(doc, 'Example: Goal 1 (30%) + Goal 2 (25%) + Goal 3 (20%) + Goal 4 (15%) + Goal 5 (10%) = 100%')

    heading(doc, 'Goal Approval Flow', level=2)
    add_flow_box(doc, '', [
        'Employee creates a goal \u2192 Status: DRAFT (can be freely edited)',
        'Employee submits for approval \u2192 Status: PENDING APPROVAL',
        'Manager reviews: APPROVE or REJECT (with reason)',
        'If rejected \u2192 Employee revises and resubmits (back to step 2)',
        'If approved \u2192 Goal is locked in. Repeat until all 5 are approved.',
    ])

    heading(doc, 'Who Can Create Goals', level=2)
    bullet(doc, 'Employee creates their own goals \u2192 submits to manager for approval', 'Employee')
    bullet(doc, 'Manager can create goals for their team member \u2192 employee is notified', 'Manager')
    bullet(doc, 'HR can create goals for any employee', 'HR')

    heading(doc, 'Audit Trail', level=2)
    bullet(doc, 'Every time a goal is submitted or resubmitted, a version record is saved', 'version record')
    bullet(doc, 'The system tracks who created, approved, or rejected each goal, and when')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 6: COMPETENCIES
    # ════════════════════════════════════════════════════════════════
    heading(doc, '6. Competency / Attribute Ratings', level=1)
    body(doc, 'Competencies (also called "Attributes") measure HOW an employee works, not just WHAT they achieve.')

    add_flow_box(doc, 'How It Works:', [
        'HR creates competency templates for each cycle (up to 5 per cycle)\n   Example: Communication, Leadership, Problem Solving, Teamwork, Innovation\n   Can also be uploaded in bulk via Excel',
        "When an employee's appraisal is created, they automatically receive all the competencies for that cycle",
        'Employee rates themselves on each competency (1\u20135 scale) with optional comments',
        'Manager rates the employee on each competency (1\u20135 scale) with optional comments',
    ])

    heading(doc, 'Scoring', level=2)
    bullet(doc, 'The competency score = average of all manager ratings', 'average of all manager ratings')
    body(doc, 'Example: Communication (4) + Leadership (3) + Problem Solving (5) + Teamwork (4) + Innovation (3) = 3.8 average',
         ['3.8 average'])
    bullet(doc, 'This average is then weighted according to the cycle configuration (e.g., 30% of the total)')

    # ════════════════════════════════════════════════════════════════
    #  SECTION 7: PEER FEEDBACK
    # ════════════════════════════════════════════════════════════════
    heading(doc, '7. Peer Feedback', level=1)
    body(doc, 'Peer feedback allows colleagues to provide ratings and comments about an employee.')

    add_styled_table(doc,
        ['Step', 'Who', 'What Happens'],
        [
            ['1', 'Employee, Manager, or HR', 'Requests feedback from a specific colleague'],
            ['2', 'System', 'Sends the request to the colleague (status: Pending)'],
            ['3', 'Colleague', 'Provides a rating (1\u20135) and optional comments'],
            ['4', 'System', 'Marks feedback as Submitted'],
            ['5', 'System', 'Includes the rating in score calculation (if peer weight > 0%)'],
        ],
        col_widths=[2, 5, 11]
    )

    heading(doc, 'Rules', level=2)
    bullet(doc, 'An employee cannot review themselves (system prevents this)', 'cannot review themselves')
    bullet(doc, 'Only one request per colleague per appraisal (no duplicates)', 'one request per colleague')
    bullet(doc, 'Once submitted, feedback cannot be changed', 'cannot be changed')
    bullet(doc, 'A pending request can be cancelled by the employee, their manager, or HR')
    bullet(doc, 'Peer feedback is only used in the final score if the cycle has peer_feedback_weight > 0%')

    heading(doc, 'Who Can See Peer Feedback', level=2)
    bullet(doc, 'The reviewer (who wrote it)')
    bullet(doc, 'The employee being reviewed')
    bullet(doc, 'Their manager')
    bullet(doc, 'HR')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 8: SCORE CALCULATION
    # ════════════════════════════════════════════════════════════════
    heading(doc, '8. How the Final Score Is Calculated', level=1)
    body(doc, 'The final score is a weighted combination of three components:')

    add_info_box(doc,
        'Final Score = (Goal Score \u00D7 Goal Weight%) + (Competency Score \u00D7 Competency Weight%) + (Peer Score \u00D7 Peer Weight%)',
        LIGHT_BLUE_BG)

    heading(doc, 'Component 1: Goal Score (Weighted Average)', level=2)
    body(doc, 'Each performance goal has a manager rating (1\u20135) and a weight. The goal score is a weighted average:')
    add_styled_table(doc,
        ['Goal', 'Manager Rating', 'Weight', 'Contribution'],
        [
            ['Increase sales by 20%', '4', '30%', '4 \u00D7 30 = 120'],
            ['Launch new product', '3', '25%', '3 \u00D7 25 = 75'],
            ['Reduce costs by 10%', '5', '20%', '5 \u00D7 20 = 100'],
            ['Improve team NPS', '4', '15%', '4 \u00D7 15 = 60'],
            ['Complete training programme', '3', '10%', '3 \u00D7 10 = 30'],
            ['TOTAL', '', '100%', '385 \u00F7 100 = 3.85'],
        ],
        col_widths=[6, 3.5, 3, 5.5],
        header_color=LIGHT_GREEN_BG
    )

    heading(doc, 'Component 2: Competency Score (Simple Average)', level=2)
    body(doc, 'The simple average of all manager competency ratings:')
    add_styled_table(doc,
        ['Competency', 'Manager Rating'],
        [
            ['Communication', '4'],
            ['Leadership', '3'],
            ['Problem Solving', '5'],
            ['Teamwork', '4'],
            ['Innovation', '3'],
            ['AVERAGE', '3.8'],
        ],
        col_widths=[9, 9]
    )

    heading(doc, 'Component 3: Peer Feedback Score (Simple Average)', level=2)
    body(doc, 'The simple average of all submitted peer ratings (1\u20135).')

    heading(doc, 'Putting It All Together \u2014 Worked Example', level=2)
    add_styled_table(doc,
        ['Component', 'Score', 'Weight', 'Contribution'],
        [
            ['Goals', '3.85', '70%', '3.85 \u00D7 0.70 = 2.695'],
            ['Competencies', '3.80', '30%', '3.80 \u00D7 0.30 = 1.140'],
            ['Peer Feedback', '\u2014', '0%', 'Not used this cycle'],
            ['FINAL SCORE', '', '', '3.835 \u2192 Rounded to 4'],
        ],
        col_widths=[4.5, 3, 3, 7.5],
        header_color=LIGHT_ORANGE_BG
    )

    heading(doc, 'Special Case: Peer Weight Exists but No One Submitted', level=2)
    bullet(doc, 'The system does NOT penalise the employee', 'does NOT penalise')
    bullet(doc, 'Instead, it redistributes the peer weight proportionally to goals and competencies', 'redistributes')
    body(doc, 'Example: Original weights: Goals 60% + Competencies 20% + Peer 20%. '
         'If no peer feedback is submitted, the system recalculates as: Goals 75% + Competencies 25% + Peer 0%. '
         'The ratio between goals and competencies stays the same (3:1).',
         ['Goals 75% + Competencies 25% + Peer 0%'])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 9: CALIBRATION
    # ════════════════════════════════════════════════════════════════
    heading(doc, '9. Calibration (Optional)', level=1)
    body(doc, 'Calibration is an optional step that HR can enable per cycle to ensure fairness across the organisation.')

    heading(doc, 'When Is It Used?', level=2)
    bullet(doc, 'Only when HR sets "Requires Calibration = Yes" in the cycle settings', 'Requires Calibration = Yes')
    bullet(doc, 'It happens after the manager submits and before the employee sees the result', 'after the manager submits')

    add_flow_box(doc, 'Calibration Flow:', [
        'Manager submits their review',
        'Instead of going to the employee, the appraisal goes to HR',
        'HR reviews all ratings across departments for consistency and fairness',
        'HR can adjust the overall rating (1\u20135) and add calibration notes',
        'HR approves \u2192 the appraisal goes to the employee for acknowledgement',
    ])

    heading(doc, 'Key Points', level=2)
    bullet(doc, 'HR can trigger calibration for all appraisals at once (bulk action)', 'all appraisals at once')
    bullet(doc, 'HR can change the overall rating (the number the employee sees)', 'change the overall rating')
    bullet(doc, 'The original calculated score is still preserved in the system for reference', 'still preserved')
    bullet(doc, 'Calibration notes are appended to the overall comments')

    # ════════════════════════════════════════════════════════════════
    #  SECTION 10: ACKNOWLEDGEMENT
    # ════════════════════════════════════════════════════════════════
    heading(doc, '10. Employee Acknowledgement & Disputes', level=1)

    add_styled_table(doc,
        ['Action', 'Details'],
        [
            ['Employee sees their final rating', 'All scores, manager comments, strengths, development areas'],
            ['Employee must acknowledge', 'This confirms they have read the appraisal'],
            ['Employee can add comments', 'Optional \u2014 any final thoughts'],
            ['Employee can flag a dispute', 'Ticking "I disagree" \u2014 the manager is automatically notified'],
        ],
        col_widths=[5.5, 12.5]
    )

    add_info_box(doc,
        'IMPORTANT: Acknowledging does NOT mean agreeing \u2014 it means "I have read and received this". '
        'The appraisal still moves to Completed after acknowledgement, even if disputed. '
        'After completion, the employee can still raise a formal Appeal.',
        LIGHT_ORANGE_BG, icon="\u26A0")

    # ════════════════════════════════════════════════════════════════
    #  SECTION 11: APPEALS
    # ════════════════════════════════════════════════════════════════
    heading(doc, '11. Appeals (After Completion)', level=1)
    body(doc, 'An appeal is a formal process for an employee to challenge their completed appraisal.')

    add_flow_box(doc, 'Appeal Process:', [
        'Employee raises appeal with a written reason',
        'HR receives the appeal (status: Pending)',
        'HR investigates (status: Under Review)',
        'HR decides: UPHELD (rating stays the same) or OVERTURNED (rating changed to a new value)',
        'Resolution and any changes are recorded in the appraisal comments',
    ])

    heading(doc, 'Rules', level=2)
    bullet(doc, 'Only the employee can raise an appeal', 'employee')
    bullet(doc, 'Only on completed appraisals', 'completed')
    bullet(doc, 'One appeal per appraisal (cannot appeal twice)', 'One appeal')
    bullet(doc, 'Must provide a written reason (cannot be empty)', 'written reason')
    bullet(doc, 'If overturned, HR provides a new overall rating (1\u20135)', 'overturned')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 12: PROBATION
    # ════════════════════════════════════════════════════════════════
    heading(doc, '12. Probation Employees (Special Handling)', level=1)
    body(doc, 'New employees who are still within their probation period get special treatment.')

    heading(doc, 'The 3-Month Rule', level=2)
    bullet(doc, 'Every new employee has a mandatory 3-month probation period', '3-month probation')
    bullet(doc, 'During this time, they are NOT eligible for the annual appraisal cycle', 'NOT eligible')
    bullet(doc, 'Instead, they are placed in a separate Probation Cycle', 'Probation Cycle')

    heading(doc, 'Automatic Handling', level=2)
    add_styled_table(doc,
        ['Employee Joined...', 'Placed In...'],
        [
            ['Before the eligibility cutoff date', 'Annual Cycle (normal)'],
            ['After the eligibility cutoff date', 'Probation Cycle (created automatically)'],
        ],
        col_widths=[9, 9],
        header_color=LIGHT_GREEN_BG
    )

    heading(doc, 'What HR Needs to Know', level=2)
    bullet(doc, 'The system automatically creates the probation cycle when the annual cycle is activated', 'automatically creates')
    bullet(doc, 'It is called "Probation Goals \u2014 {Year}"')
    bullet(doc, 'It is linked to the parent annual cycle')
    bullet(doc, 'All deferred employees are automatically placed in it')
    bullet(doc, 'The probation cycle follows the same workflow (goals \u2192 self-assessment \u2192 manager review \u2192 acknowledgement)')

    # ════════════════════════════════════════════════════════════════
    #  SECTION 13: MID-YEAR
    # ════════════════════════════════════════════════════════════════
    heading(doc, '13. Mid-Year Reviews (Team Transfers)', level=1)
    body(doc, 'Mid-year cycles are specifically for employees who changed teams during a cycle period.')

    heading(doc, 'How It Works', level=2)
    bullet(doc, 'HR creates a mid-year cycle with a date range')
    bullet(doc, 'The system checks if the employee had a team transfer during that period', 'team transfer')
    bullet(doc, 'If yes \u2192 they are eligible for the mid-year review')
    bullet(doc, 'If no transfer \u2192 they are NOT eligible')

    heading(doc, 'Why This Exists', level=2)
    body(doc, 'When someone moves to a new manager mid-cycle, neither the old nor new manager has a full picture. '
         'The mid-year review gives the new manager a structured way to evaluate the transition.')

    # ════════════════════════════════════════════════════════════════
    #  SECTION 14: LOGIN
    # ════════════════════════════════════════════════════════════════
    heading(doc, '14. Login & Access', level=1)

    add_styled_table(doc,
        ['Method', 'How It Works'],
        [
            ['Azure AD (SSO)', 'Employees click "Sign in with Microsoft" \u2014 uses the company\'s Azure Active Directory. No separate password needed.'],
            ['Local Account', 'Email and password (for users not on Azure AD)'],
        ],
        col_widths=[4.5, 13.5]
    )

    heading(doc, 'Azure AD Special Behaviours', level=2)
    bullet(doc, "The system reads the employee's role from Azure AD (Super Admin, HR Admin, Manager, Employee)", 'role from Azure AD')
    bullet(doc, 'If Azure shows someone has direct reports, they are automatically set as a Manager', 'automatically set as a Manager')
    bullet(doc, 'Department, job title, and manager info are synced from Azure AD', 'synced from Azure AD')
    bullet(doc, 'If someone transfers teams, it is recorded for mid-year review eligibility', 'transfers teams')

    heading(doc, 'Security', level=2)
    bullet(doc, 'Local accounts: After 5 failed login attempts, the account is locked for 15 minutes', '5 failed login attempts')
    bullet(doc, 'Azure AD accounts: Lockout is handled by Microsoft (not the app)')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SECTION 15: VALIDATION CHECKLIST
    # ════════════════════════════════════════════════════════════════
    heading(doc, '15. Summary of Key Rules \u2014 Validation Checklist', level=1)
    body(doc, 'Please review each rule below and mark whether it is correct or needs to be changed:')
    doc.add_paragraph('')

    rules = [
        ['1', 'Performance goals required per employee', 'Exactly 5', ''],
        ['2', 'Performance goal weights', 'Must sum to exactly 100%', ''],
        ['3', 'Development goals allowed', 'Up to 3 (optional, no score impact)', ''],
        ['4', 'Competency templates per cycle', 'Up to 5', ''],
        ['5', 'Probation period', '3 months mandatory before annual eligibility', ''],
        ['6', 'Default score weights', 'Goals 70% + Competencies 30% + Peer 0%', ''],
        ['7', 'Peer feedback with no submissions', 'Weight redistributed (employee NOT penalised)', ''],
        ['8', 'Calibration', 'Optional per cycle, HR can override final rating', ''],
        ['9', 'Employee acknowledgement', 'Required, but does NOT mean agreement', ''],
        ['10', 'Disputes', 'Flagged during acknowledgement, manager notified', ''],
        ['11', 'Appeals', 'One per appraisal, post-completion only, HR decides', ''],
        ['12', 'Active cycles', 'Only ONE per type at any time', ''],
        ['13', 'Rating scale', '1 to 5 for all ratings (goals, competencies, peer)', ''],
        ['14', 'Account lockout', '5 failed attempts = 15-minute lock (local accounts)', ''],
        ['15', 'Self-reviews in peer feedback', 'Not allowed (system prevents it)', ''],
        ['16', 'Goal approval required by', 'Manager (or HR)', ''],
        ['17', 'Who can raise an appeal', 'Employee only', ''],
        ['18', 'Probation cycle creation', 'Automatic when annual is activated', ''],
        ['19', 'Manager auto-detection', 'Via Azure AD direct reports', ''],
        ['20', 'Deadlines enforced for', 'Self-assessment and Manager review', ''],
    ]

    checklist_table = doc.add_table(rows=1 + len(rules), cols=5)
    checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    checklist_table.style = 'Table Grid'

    # Header
    headers_list = ['#', 'Rule', 'Current Behaviour', 'Correct?\n(\u2713 or \u2717)', 'HR Comments']
    widths = [1, 5.5, 5.5, 2.5, 3.5]
    for i, h in enumerate(headers_list):
        cell = checklist_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '2B579A')

    # Data
    for r_idx, row_data in enumerate(rules):
        for c_idx, val in enumerate(row_data):
            cell = checklist_table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            run.font.color.rgb = BRAND_LIGHT
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY_BG)
        # Empty columns for HR to fill
        for fill_col in [3, 4]:
            cell = checklist_table.rows[r_idx + 1].cells[fill_col]
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY_BG)

    for row in checklist_table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)

    doc.add_paragraph('')
    doc.add_paragraph('')

    # ════════════════════════════════════════════════════════════════
    #  FEEDBACK SECTION
    # ════════════════════════════════════════════════════════════════
    heading(doc, 'How to Provide Your Feedback', level=1)
    body(doc, 'For each section, please note:')
    bullet(doc, 'Is the process correct? Does it match what you expect to happen?')
    bullet(doc, 'Are the numbers right? (5 goals, 3 dev goals, 5 competencies, 3-month probation, etc.)')
    bullet(doc, 'Are the weights right? (default 70/30/0 split)')
    bullet(doc, 'Is anything missing? Any step or rule that should exist but does not?')
    bullet(doc, 'Is anything wrong? Any step that should work differently?')
    doc.add_paragraph('')
    body(doc, 'Please send your feedback with section numbers so we can address each point precisely.')

    # Footer
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2014 End of Document \u2014')
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_LIGHT
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Document generated: March 2026  |  EAS v1.0  |  CONFIDENTIAL')
    run.font.size = Pt(8)
    run.font.color.rgb = BRAND_LIGHT

    return doc


if __name__ == '__main__':
    output_dir = os.path.dirname(os.path.abspath(__file__))

    # ── Generate Word Document ──
    print("Generating Word document...")
    doc = build_document()
    docx_path = os.path.join(output_dir, 'EAS_Business_Logic_HR_Review.docx')
    doc.save(docx_path)
    print(f"  Word: {docx_path}")

    # ── Generate PDF ──
    print("Generating PDF document...")
    try:
        from fpdf import FPDF
        import re

        class EASPDF(FPDF):
            def header(self):
                if self.page_no() > 1:
                    self.set_font('Helvetica', 'I', 8)
                    self.set_text_color(91, 91, 91)
                    self.cell(0, 8, 'EAS Business Logic & Process Flow | CONFIDENTIAL', align='L')
                    self.cell(0, 8, f'Page {self.page_no()}', align='R', new_x="LMARGIN", new_y="NEXT")
                    self.set_draw_color(43, 87, 154)
                    self.line(10, 14, 200, 14)
                    self.ln(4)

            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, 'March 2026 | EAS v1.0', align='C')

            def section_heading(self, text):
                self.ln(4)
                self.set_font('Helvetica', 'B', 14)
                self.set_text_color(27, 38, 59)
                self.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
                self.set_draw_color(43, 87, 154)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(3)

            def sub_heading(self, text):
                self.ln(2)
                self.set_font('Helvetica', 'B', 11)
                self.set_text_color(43, 87, 154)
                self.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
                self.ln(1)

            def body_text(self, text):
                self.set_font('Helvetica', '', 10)
                self.set_text_color(91, 91, 91)
                self.multi_cell(0, 5.5, text)
                self.ln(2)

            def bullet_text(self, text):
                self.set_font('Helvetica', '', 9.5)
                self.set_text_color(91, 91, 91)
                self.cell(6, 5.5, '-')
                self.multi_cell(0, 5.5, text)
                self.ln(1)

            def info_box(self, text, r=253, g=232, b=208):
                self.set_fill_color(r, g, b)
                self.set_font('Helvetica', 'I', 9)
                self.set_text_color(27, 38, 59)
                self.multi_cell(0, 5.5, text, fill=True, border=1)
                self.ln(3)

            def add_table(self, headers, rows, col_widths=None):
                if not col_widths:
                    w = 190 / len(headers)
                    col_widths = [w] * len(headers)

                # Header
                self.set_font('Helvetica', 'B', 8.5)
                self.set_fill_color(43, 87, 154)
                self.set_text_color(255, 255, 255)
                for i, h in enumerate(headers):
                    self.cell(col_widths[i], 7, h, border=1, fill=True, align='L')
                self.ln()

                # Rows
                self.set_font('Helvetica', '', 8.5)
                for r_idx, row in enumerate(rows):
                    if r_idx % 2 == 1:
                        self.set_fill_color(242, 242, 242)
                    else:
                        self.set_fill_color(255, 255, 255)
                    self.set_text_color(91, 91, 91)

                    # Calculate max height needed
                    max_lines = 1
                    for c_idx, cell in enumerate(row):
                        lines = self.multi_cell(col_widths[c_idx], 5, str(cell), dry_run=True, output="LINES")
                        max_lines = max(max_lines, len(lines))
                    row_h = max(7, max_lines * 5)

                    for c_idx, cell in enumerate(row):
                        self.cell(col_widths[c_idx], row_h, str(cell)[:80], border=1, fill=True, align='L')
                    self.ln()
                self.ln(3)

        pdf = EASPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_margins(15, 15, 15)

        # Cover page
        pdf.add_page()
        pdf.ln(60)
        pdf.set_font('Helvetica', 'B', 28)
        pdf.set_text_color(27, 38, 59)
        pdf.cell(0, 15, 'Employee Appraisal System', align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.set_font('Helvetica', '', 16)
        pdf.set_text_color(43, 87, 154)
        pdf.cell(0, 12, 'Business Logic & Process Flow', align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)
        pdf.set_font('Helvetica', '', 13)
        pdf.set_text_color(23, 132, 78)
        pdf.cell(0, 10, 'For HR Team Review & Validation', align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(30)
        pdf.set_font('Helvetica', 'B', 11)
        pdf.set_text_color(232, 108, 0)
        pdf.cell(0, 10, 'CONFIDENTIAL', align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(91, 91, 91)
        pdf.cell(0, 8, 'March 2026  |  Version 1.0', align='C', new_x="LMARGIN", new_y="NEXT")

        # Content pages
        pdf.add_page()
        pdf.section_heading('1. Who Can Do What (Roles)')
        pdf.body_text('The system has four roles. Each person sees only what is relevant to them.')
        pdf.add_table(
            ['Role', 'What They Can Do'],
            [
                ['Employee', 'Set goals, self-assess, rate competencies, acknowledge appraisal, raise appeals'],
                ['Manager', 'Approve/reject goals, review & rate team, submit assessments, request peer feedback'],
                ['HR Admin', 'Manage cycles, set up templates, calibrate, handle appeals, view all appraisals'],
                ['Super Admin', 'Everything HR can do + system settings & user management'],
            ],
            [35, 155]
        )

        pdf.section_heading('2. Appraisal Cycle Setup')
        pdf.body_text('An appraisal cycle is the time period during which appraisals happen.')
        pdf.add_table(
            ['Setting', 'What It Means'],
            [
                ['Cycle Type', 'Annual, Probation, or Mid-Year'],
                ['Start & End Dates', 'The period the cycle covers'],
                ['Self-Assessment Deadline', 'Last date for employees to complete self-ratings'],
                ['Manager Review Deadline', 'Last date for managers to submit reviews'],
                ['Eligibility Cutoff', 'Employees joining after this date go to probation cycle'],
                ['Min. Service Months', 'Minimum months worked to be eligible'],
                ['Include Probation?', 'Whether probation employees are included'],
                ['Allow Proration?', 'Whether mid-cycle joiners get prorated review'],
                ['Requires Calibration?', 'Whether HR calibration step is added'],
                ['Score Weights', 'Goals % + Competencies % + Peer % = 100%'],
            ],
            [50, 140]
        )

        pdf.sub_heading('Default Score Weights')
        pdf.add_table(
            ['Component', 'Weight', 'Meaning'],
            [
                ['Goals', '70%', 'How well they achieved objectives'],
                ['Competencies', '30%', 'How they demonstrated behaviours'],
                ['Peer Feedback', '0%', 'Not used by default'],
                ['TOTAL', '100%', 'Must always add up to 100%'],
            ],
            [40, 25, 125]
        )

        pdf.add_page()
        pdf.section_heading('3. Who Is Eligible for Appraisal')
        pdf.body_text('The system checks each employee against these rules (in order):')
        pdf.add_table(
            ['Check', 'If Fails', 'Result'],
            [
                ['Completed 3-month probation?', 'Still in probation', 'Deferred to Probation Cycle'],
                ['Probation ends before cycle ends?', 'Extends beyond cycle', 'Deferred to Probation Cycle'],
                ['Joined before cutoff date?', 'Joined after cutoff', 'Deferred to Probation Cycle'],
                ['Minimum service months met?', 'Insufficient service', 'Not Eligible'],
                ['Probation employees included?', 'Excluded by cycle', 'Not Eligible'],
                ['Proration allowed? (mid-cycle join)', 'Disabled', 'Not Eligible'],
                ['Joined after cycle start?', 'With proration on', 'Eligible, PRORATED'],
                ['All checks pass', '-', 'Fully Eligible'],
            ],
            [55, 45, 90]
        )

        pdf.add_page()
        pdf.section_heading('4. The Appraisal Journey')
        pdf.body_text('Every appraisal moves through these stages in order:')
        pdf.body_text('STAGE 1: Goal Setting --> STAGE 2: Self-Assessment --> STAGE 3: Manager Review --> (Calibration if enabled) --> STAGE 4: Acknowledgement --> COMPLETED')

        pdf.sub_heading('Stage 1: Goal Setting')
        pdf.add_table(
            ['What Happens', 'Who'],
            [
                ['Performance goals created (exactly 5 required)', 'Employee / Manager'],
                ['Each goal gets a weight (must sum to 100%)', 'Employee / Manager'],
                ['Goals submitted for approval', 'Employee'],
                ['Manager approves or rejects each goal', 'Manager'],
                ['If rejected, employee revises and resubmits', 'Employee'],
                ['Development goals (optional, max 3, no score impact)', 'Employee'],
            ],
            [120, 70]
        )

        pdf.sub_heading('Stage 2: Self-Assessment')
        pdf.add_table(
            ['What Happens', 'Who'],
            [
                ['Employee rates themselves on each goal (1-5)', 'Employee'],
                ['Employee rates themselves on each competency (1-5)', 'Employee'],
                ['Employee submits self-assessment', 'Employee'],
                ['Deadline enforced', 'System'],
            ],
            [120, 70]
        )

        pdf.sub_heading('Stage 3: Manager Review')
        pdf.add_table(
            ['What Happens', 'Who'],
            [
                ['Manager rates each goal (1-5)', 'Manager'],
                ['Manager rates each competency (1-5)', 'Manager'],
                ['Manager writes strengths, development areas, comments', 'Manager'],
                ['System calculates final score', 'System'],
            ],
            [120, 70]
        )

        pdf.sub_heading('Stage 4: Employee Acknowledgement')
        pdf.add_table(
            ['What Happens', 'Who'],
            [
                ['Employee sees final rating and all feedback', 'Employee'],
                ['Employee acknowledges (does NOT mean agrees)', 'Employee'],
                ['Employee can add comments or flag dispute', 'Employee'],
                ['Appraisal marked as Completed', 'System'],
            ],
            [120, 70]
        )

        pdf.add_page()
        pdf.section_heading('5. Goal Setting & Approval')
        pdf.sub_heading('Goal Types')
        pdf.add_table(
            ['Type', 'Count', 'Required?', 'Affects Score?'],
            [
                ['Performance', 'Exactly 5', 'Yes', 'Yes - weighted into final score'],
                ['Development', 'Up to 3', 'No - optional', 'No - personal growth only'],
            ],
            [35, 30, 45, 80]
        )
        pdf.body_text('Goal Approval Flow: Draft --> Pending Approval --> Approved (or Rejected --> Revise --> Resubmit)')

        pdf.section_heading('6. Competency / Attribute Ratings')
        pdf.body_text('HR creates up to 5 competency templates per cycle. Employees and managers both rate on a 1-5 scale. The final competency score is the simple average of all manager ratings.')

        pdf.section_heading('7. Peer Feedback')
        pdf.add_table(
            ['Step', 'Who', 'What Happens'],
            [
                ['1', 'Employee/Manager/HR', 'Requests feedback from a colleague'],
                ['2', 'System', 'Sends request (status: Pending)'],
                ['3', 'Colleague', 'Provides rating (1-5) and comments'],
                ['4', 'System', 'Marks as Submitted, includes in score if weight > 0%'],
            ],
            [15, 50, 125]
        )
        pdf.bullet_text('Employee cannot review themselves')
        pdf.bullet_text('One request per colleague per appraisal (no duplicates)')
        pdf.bullet_text('Once submitted, feedback cannot be changed')

        pdf.add_page()
        pdf.section_heading('8. How the Final Score Is Calculated')
        pdf.info_box('Final Score = (Goal Score x Goal Weight) + (Competency Score x Competency Weight) + (Peer Score x Peer Weight)', 214, 228, 240)

        pdf.sub_heading('Worked Example: Goal Score (Weighted Average)')
        pdf.add_table(
            ['Goal', 'Rating', 'Weight', 'Contribution'],
            [
                ['Increase sales by 20%', '4', '30%', '4 x 30 = 120'],
                ['Launch new product', '3', '25%', '3 x 25 = 75'],
                ['Reduce costs by 10%', '5', '20%', '5 x 20 = 100'],
                ['Improve team NPS', '4', '15%', '4 x 15 = 60'],
                ['Complete training', '3', '10%', '3 x 10 = 30'],
                ['TOTAL', '', '100%', '385 / 100 = 3.85'],
            ],
            [60, 20, 25, 85]
        )

        pdf.sub_heading('Worked Example: Final Score')
        pdf.add_table(
            ['Component', 'Score', 'Weight', 'Contribution'],
            [
                ['Goals', '3.85', '70%', '3.85 x 0.70 = 2.695'],
                ['Competencies', '3.80', '30%', '3.80 x 0.30 = 1.140'],
                ['Peer Feedback', '-', '0%', 'Not used'],
                ['FINAL SCORE', '', '', '3.835 --> Rounded to 4'],
            ],
            [40, 25, 25, 100]
        )

        pdf.sub_heading('Peer Weight Fallback')
        pdf.body_text('If peer weight > 0% but no peers submitted: the system does NOT penalise the employee. It redistributes the weight proportionally. Example: 60/20/20 becomes 75/25/0 (ratio preserved).')

        pdf.add_page()
        pdf.section_heading('9. Calibration (Optional)')
        pdf.body_text('Only when "Requires Calibration = Yes" is set. After manager submits, appraisal goes to HR instead of employee. HR can adjust the overall rating (1-5) and add notes. Original score is preserved.')

        pdf.section_heading('10. Acknowledgement & Disputes')
        pdf.body_text('Acknowledging does NOT mean agreeing -- it means "I have read this." Employee can flag a dispute; manager is notified. Appraisal still completes. Employee can raise a formal appeal after completion.')

        pdf.section_heading('11. Appeals (After Completion)')
        pdf.body_text('Employee only. One per appraisal. Must provide written reason. HR decides: Upheld (no change) or Overturned (new rating assigned).')

        pdf.section_heading('12. Probation Employees')
        pdf.body_text('3-month mandatory probation. Employees not yet through probation are deferred to an auto-created Probation Cycle (linked to parent annual cycle).')

        pdf.section_heading('13. Mid-Year Reviews')
        pdf.body_text('For employees who changed teams during a cycle. System checks for team transfers. If transfer exists, employee is eligible for mid-year review.')

        pdf.section_heading('14. Login & Access')
        pdf.body_text('Azure AD SSO (automatic role sync, manager detection via direct reports) or local account (email/password, locked after 5 failed attempts for 15 minutes).')

        pdf.add_page()
        pdf.section_heading('15. Validation Checklist')
        pdf.body_text('Please review each rule and mark correct or needs change:')
        rules_pdf = [
            ['1', 'Performance goals per employee', 'Exactly 5'],
            ['2', 'Goal weights', 'Must sum to 100%'],
            ['3', 'Development goals', 'Up to 3 (optional)'],
            ['4', 'Competency templates per cycle', 'Up to 5'],
            ['5', 'Probation period', '3 months mandatory'],
            ['6', 'Default score weights', '70% / 30% / 0%'],
            ['7', 'No peer submissions', 'Weight redistributed'],
            ['8', 'Calibration', 'Optional, HR override'],
            ['9', 'Acknowledgement', 'Required, not agreement'],
            ['10', 'Disputes', 'Flag at acknowledgement'],
            ['11', 'Appeals', 'One per appraisal, HR decides'],
            ['12', 'Active cycles', 'One per type'],
            ['13', 'Rating scale', '1 to 5 everywhere'],
            ['14', 'Account lockout', '5 attempts = 15 min lock'],
            ['15', 'Self peer review', 'Not allowed'],
            ['16', 'Goal approval', 'Manager or HR'],
            ['17', 'Who appeals', 'Employee only'],
            ['18', 'Probation cycle', 'Auto-created'],
            ['19', 'Manager detection', 'Azure AD reports'],
            ['20', 'Deadlines', 'Self-assessment + Manager'],
        ]
        pdf.add_table(
            ['#', 'Rule', 'Current Behaviour'],
            rules_pdf,
            [10, 55, 125]
        )

        pdf_path = os.path.join(output_dir, 'EAS_Business_Logic_HR_Review.pdf')
        pdf.output(pdf_path)
        print(f"  PDF:  {pdf_path}")

    except Exception as e:
        print(f"  PDF generation failed: {e}")
        import traceback
        traceback.print_exc()

    print("\nDone!")
